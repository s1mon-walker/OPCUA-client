
# !pip install python-docx

import matplotlib.pyplot as plt
import sys
import time
from opcua import Client
from opcua import ua
import docx
from docx.shared import Inches


class ReadLogger:
    def __init__(self, client=None, url=None):
        if client:
            self.client = client

        if url:
            self.client = Client(url)

        self.programm = "4:DC_MOTOR"  # Bezeichnung CodeSys für das Programm
        self.logger = "4:logger_Drehzahl_0"  # Bezeichnung des Loggers in CodeSys

        self.connect()
        self.get_nodes()
        self.get_values()
        self.process_data()
        self.save_data()
        self.plot()

    def connect(self, url=None):
        self.client.connect()

    def get_nodes(self):
        base_path = ["0:Objects", "2:DeviceSet", "4:CODESYS Control for Raspberry Pi SL", "3:Resources", "4:app", "3:Programs"]
        base_path += [self.programm, self.logger]
        print(base_path)
        root = self.client.get_root_node()
        self.arTime = root.get_child(base_path + ["4:arTime"])
        self.arSignal_1 = root.get_child(base_path + ["4:arSignal_1"])  #rW_geRampt
        self.arSignal_2 = root.get_child(base_path + ["4:arSignal_2"])  #rW_ungeRampt
        self.arSignal_3 = root.get_child(base_path + ["4:arSignal_3"])  #rE
        self.arSignal_4 = root.get_child(base_path + ["4:arSignal_4"])  #rY_RE
        self.arSignal_5 = root.get_child(base_path + ["4:arSignal_5"])  #rY
        self.arSignal_6 = root.get_child(base_path + ["4:arSignal_6"])  #rY_lin
        self.arSignal_7 = root.get_child(base_path + ["4:arSignal_7"])  #rMotor_PWM
        self.arSignal_8 = root.get_child(base_path + ["4:arSignal_8"])  #rW_M
        self.arSignal_9 = root.get_child(base_path + ["4:arSignal_9"])  #rX_M
        self.arSignal_10 = root.get_child(base_path + ["4:arSignal_10"])  #rY_M

    def get_values(self):
        self.arTimeVal = self.arTime.get_value()
        self.arSignal_1val = self.arSignal_1.get_value()
        self.arSignal_2val = self.arSignal_2.get_value()
        self.arSignal_3val = self.arSignal_3.get_value()
        self.arSignal_4val = self.arSignal_4.get_value()
        self.arSignal_5val = self.arSignal_5.get_value()
        self.arSignal_6val = self.arSignal_6.get_value()
        self.arSignal_7val = self.arSignal_7.get_value()
        self.arSignal_8val = self.arSignal_8.get_value()
        self.arSignal_9val = self.arSignal_9.get_value()
        self.arSignal_10val = self.arSignal_10.get_value()

    def process_data(self):
        for i in range(0, len(self.arTimeVal)):
            self.arSignal_1val[i] = self.arSignal_1val[i]*1.0 + 0.0
            self.arSignal_2val[i] = self.arSignal_2val[i]*1.0 + 0.0
            self.arSignal_3val[i] = self.arSignal_3val[i]*1.0 + 0.0
            self.arSignal_4val[i] = self.arSignal_4val[i]*1.0 + 0.0
            self.arSignal_5val[i] = self.arSignal_5val[i]*1.0 + 0.0
            self.arSignal_6val[i] = self.arSignal_6val[i]*1.0 + 0.0
            self.arSignal_7val[i] = self.arSignal_7val[i]*1.0 + 0.0
            self.arSignal_8val[i] = self.arSignal_8val[i]*1.0 + 0.0
            self.arSignal_9val[i] = self.arSignal_9val[i]*1.0 + 0.0
            self.arSignal_10val[i] = self.arSignal_10val[i]*1.0 + 0.0

    def save_data(self):
        txtFileName = "messung.txt"
        txtFile = open(txtFileName, 'w')
        for i in range(0, len(self.arTimeVal)):
            line = ''
            line = line + '{:.3f}\t'.format(self.arTimeVal[i])
            line = line + '{:.3f}\t'.format(self.arSignal_1val[i])
            line = line + '{:.3f}\t'.format(self.arSignal_2val[i])
            line = line + '{:.3f}\t'.format(self.arSignal_3val[i])
            line = line + '{:.3f}\t'.format(self.arSignal_4val[i])
            line = line + '{:.3f}\t'.format(self.arSignal_5val[i])
            line = line + '{:.3f}\t'.format(self.arSignal_6val[i])
            line = line + '{:.3f}\t'.format(self.arSignal_7val[i])
            line = line + '{:.3f}\t'.format(self.arSignal_8val[i])
            line = line + '{:.3f}\t'.format(self.arSignal_9val[i])
            line = line + '{:.3f}'.format(self.arSignal_10val[i])
            txtFile.write(line + '\r')
        txtFile.close()

    def plot(self):
        fig = plt.figure(figsize=(15.0, 10.0), dpi=300, facecolor="white")
        ax = fig.add_subplot(1,1,1)
        ax.plot(self.arTimeVal, self.arSignal_1val, '-', linewidth=2, label="w_geRampt")
        ax.plot(self.arTimeVal, self.arSignal_2val, '-', linewidth=2, label="w")
        ax.plot(self.arTimeVal, self.arSignal_3val, '-', linewidth=2, label="x") # war mal e
        ax.plot(self.arTimeVal, self.arSignal_4val, '-', linewidth=2, label="y_RE")
        ax.plot(self.arTimeVal, self.arSignal_5val, '-', linewidth=2, label="y")
        ax.plot(self.arTimeVal, self.arSignal_6val, '-', linewidth=2, label="y_lin")
        ax.plot(self.arTimeVal, self.arSignal_7val, '-', linewidth=2, label="pwm")
        ax.plot(self.arTimeVal, self.arSignal_8val, '-', linewidth=2, label="w_M")
        ax.plot(self.arTimeVal, self.arSignal_9val, '-', linewidth=2, label="x_M")
        ax.plot(self.arTimeVal, self.arSignal_10val, '-', linewidth=2, label="y_M")

        ax.legend(loc='upper right', fancybox=True, title='')
        plt.grid(True)
        plt.show()
        fig.savefig('messung.png', dpi=96)

    def generate_word_doc(self):
        #print('Report im Word-Format') siehe https://python-docx.readthedocs.io/en/latest/
        doc = docx.Document()
        doc.add_heading('Messung', 1)
        #doc.add_paragraph('Text....')
        doc.add_picture('messung.png', width=Inches(15.0/2.54))
        doc.save('messung.docx')


if __name__ == '__main__':
    url = "opc.tcp://192.168.0.31:4840"
    logger = ReadLogger(url=url)


"""
#url = "opc.tcp://127.0.0.1:4840"
#url = "opc.tcp://172.16.3.229:4840"
url = "opc.tcp://192.168.1.120:4840"
print('OPC UA Server: ', url)
client = Client(url)

programm = "4:RET_APP" # Bezeichnung CodeSys für das Programm
logger = "4:logger" # Bezeichnung des Loggers in CodeSys

try:
    client.connect() 
    root = client.get_root_node()  
    arTime = root.get_child(["0:Objects", "2:DeviceSet", "4:CODESYS Control for Raspberry Pi SL", "3:Resources", "4:app", "3:Programs", programm, logger, "4:arTime"])
    arSignal_1 = root.get_child(["0:Objects", "2:DeviceSet", "4:CODESYS Control for Raspberry Pi SL", "3:Resources", "4:app", "3:Programs", programm, logger, "4:arSignal_1"])
    arSignal_2 = root.get_child(["0:Objects", "2:DeviceSet", "4:CODESYS Control for Raspberry Pi SL", "3:Resources", "4:app", "3:Programs", programm, logger, "4:arSignal_2"])
    arSignal_3 = root.get_child(["0:Objects", "2:DeviceSet", "4:CODESYS Control for Raspberry Pi SL", "3:Resources", "4:app", "3:Programs", programm, logger, "4:arSignal_3"])
    arSignal_4 = root.get_child(["0:Objects", "2:DeviceSet", "4:CODESYS Control for Raspberry Pi SL", "3:Resources", "4:app", "3:Programs", programm, logger, "4:arSignal_4"])
    arSignal_5 = root.get_child(["0:Objects", "2:DeviceSet", "4:CODESYS Control for Raspberry Pi SL", "3:Resources", "4:app", "3:Programs", programm, logger, "4:arSignal_5"])
    arSignal_6 = root.get_child(["0:Objects", "2:DeviceSet", "4:CODESYS Control for Raspberry Pi SL", "3:Resources", "4:app", "3:Programs", programm, logger, "4:arSignal_6"])
    arSignal_7 = root.get_child(["0:Objects", "2:DeviceSet", "4:CODESYS Control for Raspberry Pi SL", "3:Resources", "4:app", "3:Programs", programm, logger, "4:arSignal_7"])
    arSignal_8 = root.get_child(["0:Objects", "2:DeviceSet", "4:CODESYS Control for Raspberry Pi SL", "3:Resources", "4:app", "3:Programs", programm, logger, "4:arSignal_8"])
    arSignal_9 = root.get_child(["0:Objects", "2:DeviceSet", "4:CODESYS Control for Raspberry Pi SL", "3:Resources", "4:app", "3:Programs", programm, logger, "4:arSignal_9"])
    arSignal_10 = root.get_child(["0:Objects", "2:DeviceSet", "4:CODESYS Control for Raspberry Pi SL", "3:Resources", "4:app", "3:Programs", programm, logger, "4:arSignal_10"])
    
    #print('Darstellung der Daten als Plot')
    arTimeVal = arTime.get_value()
    arSignal_1val = arSignal_1.get_value()
    arSignal_2val = arSignal_2.get_value()
    arSignal_3val = arSignal_3.get_value()
    arSignal_4val = arSignal_4.get_value()
    arSignal_5val = arSignal_5.get_value()
    arSignal_6val = arSignal_6.get_value()
    arSignal_7val = arSignal_7.get_value()
    arSignal_8val = arSignal_8.get_value()
    arSignal_9val = arSignal_9.get_value() 
    arSignal_10val = arSignal_10.get_value()  
    
    # Skalierung der Signalkanäle
    for i in range(0, len(arTimeVal)):
        arSignal_1val[i] = arSignal_1val[i]*1.0 + 0.0
        arSignal_2val[i] = arSignal_2val[i]*1.0 + 0.0
        arSignal_3val[i] = arSignal_3val[i]*1.0 + 0.0
        arSignal_4val[i] = arSignal_4val[i]*1.0 + 0.0
        arSignal_5val[i] = arSignal_5val[i]*1.0 + 0.0
        arSignal_6val[i] = arSignal_6val[i]*1.0 + 0.0
        arSignal_7val[i] = arSignal_7val[i]*1.0 + 0.0
        arSignal_8val[i] = arSignal_8val[i]*1.0 + 0.0
        arSignal_9val[i] = arSignal_9val[i]*1.0 + 0.0
        arSignal_10val[i] = arSignal_10val[i]*1.0 + 0.0
            
    fig = plt.figure(figsize=(15.0, 10.0), dpi=300, facecolor="white")
    ax = fig.add_subplot(1,1,1)    
    ax.plot(arTimeVal, arSignal_1val, '-', linewidth=2, label="rW")# Signal 1
    ax.plot(arTimeVal, arSignal_2val, '-', linewidth=2, label="rY_vor_Lin") # Signal 2  
    ax.plot(arTimeVal, arSignal_3val, '-', linewidth=2, label="HAL.Y") # Signal 3
    ax.plot(arTimeVal, arSignal_4val, '-', linewidth=2, label="HAL.X") # Signal 4  
    ax.plot(arTimeVal, arSignal_5val, '-', linewidth=2, label="rE") # Signal 5
    #ax.plot(arTimeVal, arSignal_6val, '-', linewidth=2, label="Reserve") # Signal 6
    #ax.plot(arTimeVal, arSignal_7val, '-', linewidth=2, label="Reserve") # Signal 7
    #ax.plot(arTimeVal, arSignal_8val, '-', linewidth=2, label="Reserve") # Signal 8
    #ax.plot(arTimeVal, arSignal_9val, '-', linewidth=2, label="Reserve") # Signal 9
    #ax.plot(arTimeVal, arSignal_10val, '-', linewidth=2, label="Reserve") # Signal 10
    
    ax.legend(loc='upper right', fancybox=True, title='')         
    plt.grid(True)
    plt.show() 
    fig.savefig('messung.png', dpi=96)
    
    #print('Speicherung der Daten im csv-Format')
    txtFileName = "messung.txt"  
    txtFile = open(txtFileName, 'w')
    for i in range(0, len(arTimeVal)):
        line = ''
        line = line + '{:.3f}\t'.format(arTimeVal[i])
        line = line + '{:.3f}\t'.format(arSignal_1val[i])
        line = line + '{:.3f}\t'.format(arSignal_2val[i])
        line = line + '{:.3f}\t'.format(arSignal_3val[i])
        line = line + '{:.3f}\t'.format(arSignal_4val[i])
        line = line + '{:.3f}\t'.format(arSignal_5val[i])
        line = line + '{:.3f}\t'.format(arSignal_6val[i])
        line = line + '{:.3f}\t'.format(arSignal_7val[i])
        line = line + '{:.3f}\t'.format(arSignal_8val[i])
        line = line + '{:.3f}\t'.format(arSignal_9val[i])
        line = line + '{:.3f}'.format(arSignal_10val[i])
        txtFile.write(line + '\r')
    txtFile.close()

    #print('Report im Word-Format') siehe https://python-docx.readthedocs.io/en/latest/
    doc = docx.Document()
    doc.add_heading('Messung', 1)
    #doc.add_paragraph('Text....')
    doc.add_picture('messung.png', width=Inches(15.0/2.54))
    doc.save('messung.docx')
    
except:
    print('Error OPC UA open read write:', sys.exc_info()[0])

finally:
    try:
        client.disconnect()
        print('Erfassung der Messung abgeschlossen')
    except:
        print('Error OPC UA disconnect', sys.exc_info()[0])


# * [messung.png](messung.png)
# * [messung.txt](messung.txt)
# * [messung.docx](messung.docx)

# In[ ]:
"""




